from __future__ import annotations

import json
import sys
from pathlib import Path

from docx import Document


def para_style_name(paragraph):
    try:
        return paragraph.style.name or ""
    except Exception:
        return ""


def main() -> None:
    docx_path = Path(sys.argv[1])
    out_path = Path(sys.argv[2])
    document = Document(docx_path)

    paragraphs = []
    headings = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = " ".join(paragraph.text.split())
        if not text:
            continue
        style = para_style_name(paragraph)
        item = {"index": index, "style": style, "text": text}
        paragraphs.append(item)
        if style.lower().startswith("heading") or text[:4].lower() in {"chươ", "mục "}:
            headings.append(item)

    tables = []
    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows[:8]:
            cells = [" ".join(cell.text.split()) for cell in row.cells[:8]]
            if any(cells):
                rows.append(cells)
        if rows:
            tables.append({"table": table_index, "rows": rows, "rowCount": len(table.rows)})

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(
        json.dumps(
            {
                "docx": str(docx_path),
                "paragraphCount": len(paragraphs),
                "tableCount": len(document.tables),
                "headings": headings,
                "paragraphs": paragraphs,
                "tables": tables,
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )


if __name__ == "__main__":
    main()
